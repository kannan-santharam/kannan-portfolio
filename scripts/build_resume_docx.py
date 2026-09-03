#!/usr/bin/env python3
"""Build the Word resumes from src/data/resumeContent.json.

Same single source of truth as scripts/build_resume.py, so the .docx and the
.pdf never drift. Layout mirrors the PDF: A4, 0.68in margins, Arial (metric
compatible with the PDF's Helvetica), navy rules under section headings.

Usage:  python3 scripts/build_resume_docx.py [dubai|india ...]
Needs:  pip install python-docx
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONTENT = json.loads((ROOT / 'src/data/resumeContent.json').read_text())

OUT = {
    'dubai': ROOT / 'Kannan_Santharam_Senior_Lead_Software_Engineer.docx',
    'india': ROOT / 'Kannan_Santharam_Senior_Lead_Software_Engineer_ind.docx',
}
assert set(OUT) == set(CONTENT['regions'])

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
BODY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x55, 0x55, 0x55)
FONT = 'Arial'
BODY_PT, SECTION_PT, NAME_PT, SUB_PT = 11.0, 11.5, 21, 9
CONTENT_IN = 6.91          # A4 width minus both margins
SEP = '  |  '


def para(doc, space_before=0, space_after=0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    return p


def run(p, text, size=BODY_PT, bold=False, italic=False, color=BODY):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    # Word needs the east-asian hint set too or it silently substitutes
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return r


def bottom_border(p, color='1F3864', size=6):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    pPr.append(borders)


def section(doc, title):
    p = para(doc, space_before=9, space_after=3)
    run(p, title, size=SECTION_PT, bold=True, color=NAVY)
    bottom_border(p)


def bullet(doc, text):
    p = para(doc, space_after=2)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.22)
    pf.first_line_indent = Inches(-0.22)
    run(p, '•\t')
    run(p, text)


def build(region):
    R = CONTENT['regions'][region]
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)
    for side in ('left_margin', 'right_margin'):
        setattr(s, side, Inches(0.68))
    s.top_margin, s.bottom_margin = Inches(0.55), Inches(0.55)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)

    core = doc.core_properties
    core.title = f"{CONTENT['identity']['name']} - {CONTENT['identity']['title']}"
    core.author = CONTENT['identity']['name']
    core.subject = R['subject']

    # ---- header
    p = para(doc, align=1)
    run(p, CONTENT['identity']['name'].upper(), size=NAME_PT, bold=True, color=BLUE)
    p = para(doc, align=1)
    run(p, CONTENT['identity']['title'] + SEP + CONTENT['identity']['tagline'], bold=True)

    c = CONTENT['contact']
    p = para(doc, align=1)
    run(p, SEP.join([R['locationLine'], R['phoneLine'], c['email']]), size=10)

    # Nationality meta row sits between the contact line and the links, matching
    # the PDF header order.
    p = para(doc, align=1)
    meta = [('Nationality:', ' ' + CONTENT['identity']['nationality'])]
    if R['visaStatus']:
        meta.append(('Visa Status:', ' ' + R['visaStatus']))
    meta.append(('Notice Period:', ' ' + CONTENT['identity']['noticePeriod']))
    for i, (label, value) in enumerate(meta):
        if i:
            run(p, SEP, size=10)
        run(p, label, size=10, bold=True)
        run(p, value, size=10)

    p = para(doc, align=1)
    run(p, SEP.join([c['linkedin'], c['github'], R['portfolio']]), size=10)

    # ---- body
    section(doc, 'PROFESSIONAL SUMMARY')
    p = para(doc, space_after=2)
    run(p, CONTENT['summaryTemplate'].format(seeking=R['seeking']))

    section(doc, 'BUSINESS IMPACT SNAPSHOT')
    for b in CONTENT['snapshot']:
        bullet(doc, b)

    section(doc, 'CORE COMPETENCIES')
    for comp in CONTENT['competencies']:
        p = para(doc, space_after=2)
        run(p, comp['label'] + ':', bold=True)
        run(p, ' ' + ', '.join(comp['items']))

    section(doc, 'PROFESSIONAL EXPERIENCE')
    for job in CONTENT['experience']:
        p = para(doc, space_before=6, space_after=0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_IN), WD_TAB_ALIGNMENT.RIGHT)
        run(p, job['title'], bold=True)
        run(p, '\t' + job['dates'], bold=True)
        p = para(doc, space_after=0)
        run(p, job['company'], bold=True, color=BLUE)
        if job['progression']:
            p = para(doc, space_after=2)
            run(p, 'Progression: ' + ' → '.join(job['progression']),
                size=SUB_PT, italic=True, color=GRAY)
        for b in job['bullets']:
            bullet(doc, b)

    section(doc, 'EDUCATION')
    ed = CONTENT['education']
    p = para(doc, space_after=0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_IN), WD_TAB_ALIGNMENT.RIGHT)
    run(p, ed['degree'], bold=True)
    run(p, '\t' + ed['years'], bold=True)
    p = para(doc, space_after=2)
    run(p, ed['institution'], color=GRAY)

    section(doc, 'LANGUAGES')
    p = para(doc, space_after=2)
    for i, lang in enumerate(CONTENT['languages']):
        if i:
            run(p, SEP)
        run(p, lang['name'] + ':', bold=True)
        run(p, ' ' + lang['level'])

    doc.save(OUT[region])
    return OUT[region]


if __name__ == '__main__':
    for region in (sys.argv[1:] or list(OUT)):
        assert region in OUT, f'unknown region {region!r}; expected {list(OUT)}'
        path = build(region)

        # read it back and prove the content survived
        d = Document(path)
        text = '\n'.join(p.text for p in d.paragraphs)
        R = CONTENT['regions'][region]
        assert R['seeking'] in text, f'{region}: seeking line missing'
        assert CONTENT['identity']['title'] in text, f'{region}: title missing'
        n_bullets = sum(1 for p in d.paragraphs if p.text.startswith('•'))
        expected = len(CONTENT['snapshot']) + sum(len(j['bullets']) for j in CONTENT['experience'])
        assert n_bullets == expected, f'{region}: {n_bullets} bullets, expected {expected}'
        for comp in CONTENT['competencies']:
            assert comp['label'] + ':' in text, f'{region}: missing competency {comp["label"]}'
        if region == 'india':
            for banned in ('Dubai', 'UAE', 'BOTIM', 'Visa Status', 'Relocate'):
                assert banned not in text, f'india: leaked {banned} in text'
                assert banned not in str(d.core_properties.subject), f'india: leaked {banned} in metadata'
        print(f'{region:6s} bullets={n_bullets} paras={len(d.paragraphs)} -> '
              f'{path.name} {path.stat().st_size} bytes')
